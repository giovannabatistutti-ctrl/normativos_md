"""
contract_reader.py (v2)
=======================
Módulo de leitura de contratos — PDF, DOCX, texto.
Padrão MCP-queue: se download do Jira falhar, salva pending para o Planner.

Funções públicas:
  - ler_contrato_anexo(anexos, ticket_id) → dict
  - ler_contrato_url(url, nome, ticket_id) → dict
  - extrair_dados_contrato(texto) → dict
  - salvar_resultado_contrato(ticket_id, dados) → Path
  - listar_contratos_pendentes() → list
"""

import json
import logging
import re
import tempfile
from datetime import datetime
from pathlib import Path

import requests

logger = logging.getLogger(__name__)

BASE_DIR = Path(__file__).parent.parent
PENDING_ATTACH_DIR = BASE_DIR / "data" / "pending_attachments"
ATTACH_RESULTS_DIR = BASE_DIR / "data" / "attachment_results"
TEMP_DIR = BASE_DIR / "data" / "temp_contracts"

# ---------------------------------------------------------------------------
# Regex para extração de campos do contrato
# ---------------------------------------------------------------------------

REGEX_CNPJ = re.compile(r"\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}")
REGEX_CEP = re.compile(r"\d{2}\.\d{3}-\d{3}|\d{5}-\d{3}")
REGEX_DATA = re.compile(r"\d{1,2}\s+de\s+\w+\s+de\s+\d{4}")
REGEX_AVISO_PREVIO = re.compile(
    r"10\.3[^\d]{0,50}(\d+)\s*\([\w\s]+\)\s*dias",
    re.IGNORECASE | re.DOTALL,
)
REGEX_ENDERECO = re.compile(
    r"(?:Rua|Avenida|Av\.|R\.|Alameda|Al\.|Travessa|Tv\.)[^\n,]{5,80},\s*n?[°º]?\s*\d+",
    re.IGNORECASE,
)
REGEX_COLABMAIS = re.compile(r"[Cc]olab\+|[Pp]rograma\s+[Cc]olab")


# ---------------------------------------------------------------------------
# Extração de texto de formatos
# ---------------------------------------------------------------------------

def _extrair_texto_pdf(caminho: Path) -> str:
    """Extrai texto de PDF usando pdfplumber."""
    try:
        import pdfplumber
        partes = []
        with pdfplumber.open(str(caminho)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    partes.append(t)
        return "\n".join(partes)
    except Exception as exc:
        logger.warning("_extrair_texto_pdf: erro — %s", exc)
        return ""


def _extrair_texto_docx(caminho: Path) -> str:
    """Extrai texto de DOCX usando python-docx."""
    try:
        from docx import Document
        doc = Document(str(caminho))
        partes = [p.text for p in doc.paragraphs if p.text.strip()]
        # Tabelas também
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        partes.append(cell.text.strip())
        return "\n".join(partes)
    except Exception as exc:
        logger.warning("_extrair_texto_docx: erro — %s", exc)
        return ""


def _extrair_texto_arquivo(caminho: Path, nome: str) -> str:
    """Detecta formato e extrai texto."""
    nome_lower = nome.lower()
    if nome_lower.endswith(".pdf"):
        return _extrair_texto_pdf(caminho)
    if nome_lower.endswith((".docx", ".doc")):
        return _extrair_texto_docx(caminho)
    # Tentar como texto puro
    try:
        return caminho.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return ""


# ---------------------------------------------------------------------------
# Extração de dados estruturados do texto do contrato
# ---------------------------------------------------------------------------

def extrair_dados_contrato(texto: str) -> dict:
    """
    Extrai campos estruturados do texto de um contrato iFood Benefícios.
    Nunca inventa — retorna None para campos não encontrados.
    """
    dados = {
        "razao_social": None,
        "cnpj": None,
        "endereco": None,
        "cep": None,
        "data_assinatura_original": None,
        "aviso_previo_clausula_10_3": None,
        "tem_colabmais": False,
    }

    # CNPJ — pegar o primeiro que NÃO é do iFood (33.157.312/0001-62)
    cnpjs = REGEX_CNPJ.findall(texto)
    cnpjs_empresa = [c for c in cnpjs if c != "33.157.312/0001-62"]
    if cnpjs_empresa:
        dados["cnpj"] = cnpjs_empresa[0]

    # CEP
    cep_match = REGEX_CEP.search(texto)
    if cep_match:
        dados["cep"] = cep_match.group()

    # Data de assinatura — pegar a primeira data mencionada
    data_matches = REGEX_DATA.findall(texto)
    if data_matches:
        dados["data_assinatura_original"] = data_matches[0]

    # Endereço
    end_match = REGEX_ENDERECO.search(texto)
    if end_match:
        dados["endereco"] = end_match.group().strip()

    # Aviso prévio (Cláusula 10.3)
    avp_match = REGEX_AVISO_PREVIO.search(texto)
    if avp_match:
        try:
            dados["aviso_previo_clausula_10_3"] = int(avp_match.group(1))
        except ValueError:
            pass

    # Colab+
    dados["tem_colabmais"] = bool(REGEX_COLABMAIS.search(texto))

    return dados


# ---------------------------------------------------------------------------
# Download de contrato
# ---------------------------------------------------------------------------

def _baixar_arquivo(url: str, nome: str) -> Path | None:
    """Tenta baixar o arquivo via HTTP. Retorna path local ou None."""
    if not url:
        return None

    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    dest = TEMP_DIR / nome

    try:
        resp = requests.get(url, timeout=30, stream=True)
    except requests.exceptions.RequestException as exc:
        logger.warning("_baixar_arquivo: erro de conexão — %s", exc)
        return None

    if resp.status_code in (401, 403, 407):
        logger.warning("_baixar_arquivo: sem permissão (HTTP %d)", resp.status_code)
        return None

    if not resp.ok:
        logger.warning("_baixar_arquivo: HTTP %d para %s", resp.status_code, url)
        return None

    # Verificar se é HTML (erro mascarado)
    ct = resp.headers.get("content-type", "")
    if "text/html" in ct:
        logger.warning("_baixar_arquivo: resposta HTML — provavelmente login page")
        return None

    with open(dest, "wb") as f:
        for chunk in resp.iter_content(chunk_size=8192):
            f.write(chunk)

    logger.info("_baixar_arquivo: %s baixado (%d bytes)", nome, dest.stat().st_size)
    return dest


# ---------------------------------------------------------------------------
# Padrão pending-file
# ---------------------------------------------------------------------------

def _salvar_pending_attachment(ticket_id: str, anexo: dict) -> Path:
    """Salva pending file para o Planner baixar o arquivo via MCP."""
    PENDING_ATTACH_DIR.mkdir(parents=True, exist_ok=True)
    slug = f"{ticket_id}_{anexo.get('nome', 'contrato')}"[:80]
    path = PENDING_ATTACH_DIR / f"{slug}_pending.json"
    path.write_text(json.dumps({
        "ticket_id": ticket_id,
        "anexo": anexo,
        "status": "aguardando_mcp",
        "criado_em": datetime.now().isoformat(timespec="seconds"),
    }, ensure_ascii=False, indent=2))
    logger.info("contract_reader: pending_attachment salvo em %s", path)
    return path


def _verificar_resultado_disponivel(ticket_id: str) -> dict | None:
    """Verifica se o Planner já processou o contrato para este ticket."""
    result_path = ATTACH_RESULTS_DIR / f"{ticket_id}_contrato.json"
    if result_path.exists():
        return json.loads(result_path.read_text(encoding="utf-8"))
    return None


def salvar_resultado_contrato(ticket_id: str, dados: dict) -> Path:
    """Chamado pelo Planner após baixar e extrair o contrato via MCP."""
    ATTACH_RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    path = ATTACH_RESULTS_DIR / f"{ticket_id}_contrato.json"
    dados["_source"] = "mcp_planner"
    dados["_processado_em"] = datetime.now().isoformat(timespec="seconds")
    path.write_text(json.dumps(dados, ensure_ascii=False, indent=2, default=str))
    return path


def listar_contratos_pendentes() -> list:
    """Lista tickets aguardando leitura de contrato pelo Planner."""
    if not PENDING_ATTACH_DIR.exists():
        return []
    pendentes = []
    for f in PENDING_ATTACH_DIR.glob("*_pending.json"):
        ticket_id = f.name.split("_")[0]
        result = ATTACH_RESULTS_DIR / f"{ticket_id}_contrato.json"
        if not result.exists():
            pendentes.append(json.loads(f.read_text(encoding="utf-8")))
    return pendentes


# ---------------------------------------------------------------------------
# Funções principais
# ---------------------------------------------------------------------------

def ler_contrato_url(url: str, nome_arquivo: str, ticket_id: str = "") -> dict:
    """
    Baixa e extrai dados de um contrato a partir de URL.
    Se falhar, salva pending para o Planner processar via MCP.
    """
    # 1. Verificar se resultado já existe
    if ticket_id:
        existente = _verificar_resultado_disponivel(ticket_id)
        if existente:
            logger.info("contract_reader: resultado existente para %s", ticket_id)
            return existente

    # 2. Tentar download direto
    arquivo_local = _baixar_arquivo(url, nome_arquivo) if url else None

    if arquivo_local:
        texto = _extrair_texto_arquivo(arquivo_local, nome_arquivo)
        dados = extrair_dados_contrato(texto)
        dados["fonte"] = "download_direto"
        dados["arquivo"] = nome_arquivo
        dados["texto_parcial"] = texto[:500] if texto else ""
        logger.info("contract_reader: contrato lido com sucesso — %s", nome_arquivo)
        return dados

    # 3. Salvar pending para o Planner
    if ticket_id and url:
        path = _salvar_pending_attachment(ticket_id, {"nome": nome_arquivo, "url": url})
        return {
            "status": "pendente_mcp",
            "ticket_id": ticket_id,
            "pending_path": str(path),
            "fonte": "pendente_mcp",
            "mensagem": f"Download direto falhou. Planner deve baixar via MCP e chamar salvar_resultado_contrato().",
        }

    return {"fonte": "sem_url", "erro": "URL de download não disponível"}


def ler_contrato_anexo(anexos: list, ticket_id: str = "") -> dict:
    """
    Função principal — tenta ler o contrato dos anexos do ticket.
    Filtra PDFs/DOCXs ignorando propostas comerciais.
    """
    if not anexos:
        logger.info("contract_reader: sem anexos para %s", ticket_id)
        return {"fonte": "sem_anexo", "razao_social": None, "cnpj": None}

    # Filtrar contratos (ignorar propostas comerciais e imagens)
    contrato_anexos = [
        a for a in anexos
        if isinstance(a, dict)
        and a.get("nome", "").lower().endswith((".pdf", ".docx", ".doc"))
        and "proposta" not in a.get("nome", "").lower()
        and "comercial" not in a.get("nome", "").lower()
    ]

    if not contrato_anexos:
        # Fallback: qualquer PDF/DOCX
        contrato_anexos = [
            a for a in anexos
            if isinstance(a, dict)
            and a.get("nome", "").lower().endswith((".pdf", ".docx", ".doc"))
        ]

    if not contrato_anexos:
        return {"fonte": "sem_contrato", "razao_social": None, "cnpj": None}

    anexo = contrato_anexos[0]
    logger.info("contract_reader: tentando ler '%s' para ticket %s", anexo.get("nome"), ticket_id)

    return ler_contrato_url(
        url=anexo.get("url", ""),
        nome_arquivo=anexo.get("nome", "contrato"),
        ticket_id=ticket_id,
    )


# ---------------------------------------------------------------------------
# Compatibilidade com código legado
# ---------------------------------------------------------------------------

class ContractReader:
    """Wrapper de classe para compatibilidade."""

    def __init__(self, config: dict = None):
        pass

    def read(self, attachment_url: str, filename: str = "contrato") -> dict:
        return ler_contrato_url(url=attachment_url, nome_arquivo=filename)
